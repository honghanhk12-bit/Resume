"""Writes a RewrittenResume into a clean, ATS-friendly .docx file.

Deliberately avoids tables, columns, text boxes, headers/footers, and
non-standard fonts/colors — all of which are known to break ATS parsers.
"""

import logging
import re

from docx import Document
from docx.shared import Pt

from schemas import RewrittenResume

logger = logging.getLogger("job_hunt_agents")

FONT_NAME = "Calibri"
BODY_SIZE = Pt(11)
HEADING_SIZE = Pt(13)
NAME_SIZE = Pt(16)


def _set_run_style(run, size=BODY_SIZE, bold=False):
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold


def _add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.upper())
    _set_run_style(run, size=HEADING_SIZE, bold=True)
    paragraph.space_after = Pt(4)


def _add_body_paragraph(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    _set_run_style(run, bold=bold)


def _add_bullet(document: Document, text: str) -> None:
    # Use a plain hyphen bullet rather than Word's native bullet list style,
    # which some ATS parsers mangle into unreadable characters.
    paragraph = document.add_paragraph(style=None)
    run = paragraph.add_run(f"- {text}")
    _set_run_style(run)


def write_resume_docx(resume: RewrittenResume, output_path: str) -> None:
    """Render a RewrittenResume to a plain, single-column .docx file.

    Args:
        resume: The final, approved RewrittenResume.
        output_path: Destination path for the .docx file.
    """
    document = Document()

    # Base style
    style = document.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE

    # Contact block (name assumed to be the first line)
    contact_lines = [line.strip() for line in resume.contact.splitlines() if line.strip()]
    if contact_lines:
        name_paragraph = document.add_paragraph()
        name_run = name_paragraph.add_run(contact_lines[0])
        _set_run_style(name_run, size=NAME_SIZE, bold=True)
        for line in contact_lines[1:]:
            _add_body_paragraph(document, line)

    # Summary
    if resume.summary:
        _add_heading(document, "Summary")
        _add_body_paragraph(document, resume.summary)

    # Experience
    if resume.experience:
        _add_heading(document, "Experience")
        for entry in resume.experience:
            header = f"{entry.title} | {entry.company} | {entry.dates}"
            _add_body_paragraph(document, header, bold=True)
            for bullet in entry.bullets:
                _add_bullet(document, _strip_ats_unsafe_chars(bullet))

    # Education
    if resume.education:
        _add_heading(document, "Education")
        for line in resume.education:
            _add_body_paragraph(document, line)

    # Skills
    if resume.skills:
        _add_heading(document, "Skills")
        _add_body_paragraph(document, " | ".join(resume.skills))

    document.save(output_path)
    logger.info("Wrote resume docx to %s", output_path)


def _strip_ats_unsafe_chars(text: str) -> str:
    """Remove characters known to confuse ATS parsers, keeping hyphens/pipes."""
    # Keep alphanumerics, standard punctuation, hyphens, and pipes.
    cleaned = re.sub(r"[^\w\s.,%$#&()/\-|+':]", "", text)
    return re.sub(r"\s+", " ", cleaned).strip()
