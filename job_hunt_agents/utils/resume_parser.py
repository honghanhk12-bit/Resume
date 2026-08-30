"""Reads a resume (PDF or DOCX) into a single clean text block."""

import logging
import os

import pdfplumber
from docx import Document

logger = logging.getLogger("job_hunt_agents")


def parse_resume(path: str) -> str:
    """Parse a resume file into plain text.

    Args:
        path: Path to a .pdf or .docx resume file.

    Returns:
        The resume's text content, with pages/paragraphs joined by newlines.

    Raises:
        FileNotFoundError: If the path does not exist.
        ValueError: If the file extension is unsupported.
    """
    if not os.path.exists(path):
        raise FileNotFoundError(
            f"Resume file not found at '{path}'. Drop your resume there or "
            f"update RESUME_PATH in config.py."
        )

    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        text = _parse_pdf(path)
    elif ext == ".docx":
        text = _parse_docx(path)
    else:
        raise ValueError(f"Unsupported resume format '{ext}'. Use .pdf or .docx.")

    cleaned = _clean_text(text)
    logger.info("Parsed resume from %s (%d characters)", path, len(cleaned))
    return cleaned


def _parse_pdf(path: str) -> str:
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            pages.append(page_text)
    return "\n".join(pages)


def _parse_docx(path: str) -> str:
    document = Document(path)
    paragraphs = [p.text for p in document.paragraphs]
    # Also pull text out of any tables, in case the resume uses them.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def _clean_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines)
